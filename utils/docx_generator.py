"""
MALT IAC Report Word Document Generator
Produces a MALT-styled .docx using python-docx.

Structure:
  Cover Page
  Disclaimer
  Preface
  Table of Contents
  Section 1: Executive Summary (utility table, AR summary table)
  Section 2: General Facility Background
  Section 3: Energy Assessment Recommendations (one subsection per AR)
  Section 4: Cybersecurity Assessment
"""

import io
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants ─────────────────────────────────────────────────────────────────

MALT_NAVY       = RGBColor(0x00, 0x33, 0x66)
MALT_LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF0)
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)
BLACK           = RGBColor(0x00, 0x00, 0x00)
HEADER_GRAY     = RGBColor(0x80, 0x80, 0x80)

DISCLAIMER_TEXT = (
    "This report was prepared as an account of work sponsored by the United States "
    "Department of Energy (DOE) Office of Manufacturing and Energy Supply Chains (MESC) "
    "through the Industrial Training and Assessment Centers (ITACs) Program and conducted "
    "by the University of Louisiana at Lafayette (UL Lafayette). Neither the DOE MESC nor "
    "UL Lafayette makes any warranty, express or implied, or assumes any legal liability or "
    "responsibility for the accuracy, completeness, or usefulness of any information, "
    "apparatus, product, or process disclosed, or represents that its use would not infringe "
    "privately owned rights."
)

PREFACE_TEXT = (
    "The work in this report is a service of the University of Louisiana at Lafayette "
    "Industrial and Building Assessment Center. The Center is financially supported by "
    "multiple sponsors, including the U.S. Department of Energy, U.S. Department of "
    "Agriculture, and Cleco. The primary objective of the Center is to identify cost savings "
    "by evaluating opportunities for energy conservation, waste minimization, and productivity "
    "enhancements for industrial facilities and commercial buildings."
)

CYBER_TEXT = (
    "As systems to control energy-using manufacturing equipment become more connected to the "
    "internet, it is important for plant operations staff to have an understanding of "
    "cybersecurity risks and to coordinate risk management activities within their organization. "
    "Small businesses may not consider themselves targets for cyber-attacks. However, they have "
    "valuable information that cyber criminals seek, such as employee and customer records, bank "
    "account information, and access to larger networks. They can be at a higher risk of a "
    "cybersecurity attack because they have fewer resources dedicated to cybersecurity. "
    "Facilities are encouraged to implement network segmentation between operational technology "
    "(OT) and information technology (IT) networks; regularly update firmware and software on "
    "all networked devices; establish and enforce strong password policies for all control "
    "systems; conduct regular cybersecurity risk assessments in accordance with the NIST "
    "Cybersecurity Framework; and develop and exercise incident response plans for cybersecurity "
    "events. The MALT IAC recommends consulting with a qualified cybersecurity professional to "
    "perform a thorough assessment of all networked systems at this facility."
)


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _hex_to_rgb_str(rgb_color: RGBColor) -> str:
    """Convert RGBColor to 6-char hex string (no #)."""
    return f"{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"


def _set_cell_bg(cell, rgb_color: RGBColor):
    """Set table cell background color via XML shading element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex_to_rgb_str(rgb_color))
    # Remove existing shd if present
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_run_font(run, name: str = "Calibri", size_pt: float = 11,
                  bold: bool = False, color: RGBColor = None):
    """Apply font styling to a run."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Ensure Calibri works in East Asian context too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _set_para_space(para, space_before_pt: float = 0, space_after_pt: float = 0):
    """Set paragraph spacing."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if space_before_pt:
        spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
    if space_after_pt:
        spacing.set(qn("w:after"), str(int(space_after_pt * 20)))


def _add_page_break(doc: Document):
    """Insert an explicit page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_col_width(table, col_idx: int, width_inches: float):
    """Set a column width in a table."""
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(width_inches * 1440)))
        tcW.set(qn("w:type"), "dxa")


def _add_running_header(doc: Document, director: str, report_no: str):
    """Add running header to all sections: right-aligned, small gray text."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header for non-first pages
    header = section.header
    # Clear existing
    for para in header.paragraphs:
        for run in para.runs:
            run.text = ""

    if header.paragraphs:
        hpara = header.paragraphs[0]
    else:
        hpara = header.add_paragraph()

    hpara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hpara.add_run(f"MALT IAC Director: {director}    Report No: {report_no}")
    _set_run_font(run, "Calibri", 8.5, False, HEADER_GRAY)


# ── Heading helpers ───────────────────────────────────────────────────────────

def _set_heading(doc: Document, text: str, level: int,
                 color: RGBColor = None, page_break_before: bool = False) -> None:
    """Add a styled heading paragraph."""
    if color is None:
        color = MALT_NAVY

    para = doc.add_paragraph()
    if page_break_before:
        pPr = para._p.get_or_add_pPr()
        pb = OxmlElement("w:pageBreakBefore")
        pb.set(qn("w:val"), "1")
        pPr.append(pb)

    run = para.add_run(text)
    if level == 1:
        _set_run_font(run, "Calibri", 16, True, color)
        _set_para_space(para, 6, 12)
    elif level == 2:
        _set_run_font(run, "Calibri", 13, True, color)
        _set_para_space(para, 4, 6)
    elif level == 3:
        _set_run_font(run, "Calibri", 11, True, color)
        _set_para_space(para, 2, 4)
    return para


def _add_body_para(doc: Document, text: str, bold: bool = False,
                   italic: bool = False, color: RGBColor = None) -> None:
    """Add a styled body paragraph."""
    if color is None:
        color = BLACK
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, "Calibri", 11, bold, color)
    if italic:
        run.font.italic = True
    _set_para_space(para, 0, 4)


# ── Table helper ──────────────────────────────────────────────────────────────

def _add_table(doc: Document, data: list, headers: list = None,
               col_widths: list = None) -> None:
    """
    Add a formatted MALT-styled table.

    data      — list of row dicts or row lists (not including header row)
    headers   — list of header strings (if None, first row of data is header)
    col_widths— list of floats (inches) for each column
    """
    if headers is not None:
        all_rows = [headers] + [list(r.values()) if isinstance(r, dict) else r for r in data]
    else:
        all_rows = [list(r.values()) if isinstance(r, dict) else r for r in data]

    if not all_rows:
        return

    num_cols = len(all_rows[0])
    num_rows = len(all_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(all_rows):
        row = table.rows[r_idx]
        for c_idx, cell_val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(cell_val) if cell_val is not None else "")
            if r_idx == 0:
                # Header row: navy bg, white bold text
                _set_cell_bg(cell, MALT_NAVY)
                _set_run_font(run, "Calibri", 10, True, WHITE)
            elif r_idx % 2 == 0:
                # Even data rows: light gray
                _set_cell_bg(cell, MALT_LIGHT_GRAY)
                _set_run_font(run, "Calibri", 10, False, BLACK)
            else:
                # Odd data rows: white
                _set_run_font(run, "Calibri", 10, False, BLACK)

            # Compact padding via XML
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side in ["top", "bottom", "left", "right"]:
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "60")
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            # Remove old margins if present
            old = tcPr.find(qn("w:tcMar"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcMar)

    # Apply column widths
    if col_widths:
        for c_idx, width in enumerate(col_widths):
            if c_idx < num_cols:
                _set_col_width(table, c_idx, width)

    doc.add_paragraph()  # spacer after table


# ── AR section helper ─────────────────────────────────────────────────────────

def _add_ar_section(doc: Document, ar: dict, ar_num: int) -> None:
    """Add a complete AR section to the document."""
    ar_number  = ar.get("ar_number", f"AR-{ar_num}")
    arc_code   = ar.get("arc_code", "")
    title      = ar.get("title", "")
    cost_sav   = ar.get("total_cost_savings", 0)
    impl_cost  = ar.get("implementation_cost", 0)
    payback    = ar.get("payback", float("inf"))

    # AR header line
    heading_text = f"{ar_number}: ARC {arc_code} — {title}"
    _set_heading(doc, heading_text, 2)

    # Resource savings summary
    res_list = ar.get("resources", [])
    res_parts = [
        f"{r['type']}: {r['savings']:,.0f} {r['unit']}"
        for r in res_list if r.get("savings", 0) > 0
    ]
    payback_str = f"{payback:.1f} yr" if payback not in (float("inf"), None) else "N/A"

    summary_headers = ["AR #", "ARC Code", "Resource Savings", "Annual Cost Savings", "Impl. Cost", "Payback"]
    summary_data = [[
        ar_number,
        arc_code,
        "\n".join(res_parts) or "—",
        f"${cost_sav:,.0f}/yr",
        f"${impl_cost:,.0f}",
        payback_str,
    ]]
    _add_table(doc, summary_data, headers=summary_headers,
               col_widths=[0.75, 0.85, 1.7, 1.4, 1.0, 0.75])

    # Observation
    if ar.get("observation"):
        _set_heading(doc, "Observation", 3)
        _add_body_para(doc, ar["observation"])

    # Recommendation
    if ar.get("recommendation"):
        _set_heading(doc, "Recommendation", 3)
        _add_body_para(doc, ar["recommendation"])

    # Technology Description
    if ar.get("tech_description"):
        _set_heading(doc, "Technology Description", 3)
        _add_body_para(doc, ar["tech_description"])

    # Calculation
    calc = ar.get("calculation_details", {})
    if calc:
        _set_heading(doc, "Calculation", 3)
        narrative = calc.get("narrative", "")
        if narrative:
            _add_body_para(doc, narrative)
        else:
            # Format key-value pairs
            skip_keys = {"narrative"}
            for key, val in calc.items():
                if key in skip_keys:
                    continue
                label = key.replace("_", " ").title()
                if isinstance(val, float):
                    val_str = f"{val:,.4f}" if abs(val) < 100 else f"{val:,.2f}"
                else:
                    val_str = str(val)
                _add_body_para(doc, f"{label}: {val_str}")

    # Divider
    para = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    para._p.get_or_add_pPr().append(pBdr)


# ── Table of Contents ─────────────────────────────────────────────────────────

def _add_toc(doc: Document, ar_list: list) -> None:
    """Add a simple text-based Table of Contents."""
    _set_heading(doc, "Table of Contents", 1)

    sections = [
        ("Disclaimer",                           ""),
        ("Preface",                              ""),
        ("1.  Executive Summary",                ""),
        ("    1.1  Annual Utility Usage and Cost",""),
        ("    1.2  Summary of Assessment Recommendations", ""),
        ("2.  General Facility Background",       ""),
        ("    2.1  Facility Description",         ""),
        ("    2.2  Process Description",          ""),
        ("    2.3  Best Practices",               ""),
        ("    2.4  Forms of Energy Usage",        ""),
        ("    2.5  Major Energy Consuming Equipment",""),
        ("    2.6  Energy and Water Consumption with Cost",""),
        ("3.  Energy Assessment Recommendations", ""),
    ]
    for idx, ar in enumerate(ar_list):
        ar_num  = ar.get("ar_number", f"AR-{idx+1}")
        arc     = ar.get("arc_code", "")
        title   = ar.get("title", "")
        sections.append((f"    3.{idx+1}  {ar_num} — {title} (ARC {arc})", ""))

    sections.append(("4.  Cybersecurity Assessment", ""))

    for text, _ in sections:
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_run_font(run, "Calibri", 11, False, BLACK)
        _set_para_space(para, 0, 2)

    _add_page_break(doc)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_docx_report(session: dict) -> bytes:
    """
    Generate full MALT Word report from session dict.
    Returns bytes that can be downloaded via st.download_button.
    """
    # Import here to avoid circular issues when running outside Streamlit
    try:
        from utils.session import get_utility_rates_from_dict
    except ImportError:
        import sys, os
        sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
        from utils.session import get_utility_rates_from_dict

    doc = Document()

    # ── Page setup: Letter, 1" margins ────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1.1)
    section.bottom_margin = Inches(0.9)

    # Session values
    report_no    = session.get("report_number", "IAC-XXXX")
    director     = session.get("lead_faculty",  'Dr. Peng "Solomon" Yin')
    visit_date   = session.get("site_visit_date", None)
    location     = session.get("location", "")
    naics        = session.get("naics_code", "")
    sic          = session.get("sic_code", "")
    principal    = session.get("principal_products", "")
    lead_student = session.get("lead_student", "")
    safety_stu   = session.get("safety_student", "")
    other_stu    = session.get("other_students", "")
    ar_list      = session.get("ar_list", [])
    include_cyber = session.get("include_cyber", True)
    include_toc   = session.get("include_toc", True)

    # Running header (non-first pages)
    _add_running_header(doc, director, report_no)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # MALT logo placeholder
    cover_logo = doc.add_paragraph()
    cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = cover_logo.add_run("MALT Industrial Assessment Center")
    _set_run_font(logo_run, "Calibri", 22, True, MALT_NAVY)
    _set_para_space(cover_logo, 36, 6)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Energy Assessment Report")
    _set_run_font(sub_run, "Calibri", 14, False, MALT_NAVY)
    _set_para_space(subtitle_para, 0, 24)

    # Cover info table
    visit_date_str = str(visit_date) if visit_date else "—"
    cover_rows = [
        ["Report Number:",     report_no],
        ["Site Visit Date:",   visit_date_str],
        ["Facility Location:", location or "—"],
        ["Principal Products:", principal or "—"],
        ["NAICS Code:",        naics or "—"],
        ["SIC Code:",          sic or "—"],
    ]
    cover_table = doc.add_table(rows=len(cover_rows), cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (label, value) in enumerate(cover_rows):
        row = cover_table.rows[r_idx]
        row.cells[0].text = ""
        label_run = row.cells[0].paragraphs[0].add_run(label)
        _set_run_font(label_run, "Calibri", 11, True, BLACK)
        row.cells[1].text = ""
        val_run = row.cells[1].paragraphs[0].add_run(value)
        _set_run_font(val_run, "Calibri", 11, False, BLACK)
    _set_col_width(cover_table, 0, 2.2)
    _set_col_width(cover_table, 1, 4.3)
    doc.add_paragraph()

    # Assessment team table
    team_label = doc.add_paragraph()
    team_label_run = team_label.add_run("Assessment Team")
    _set_run_font(team_label_run, "Calibri", 12, True, MALT_NAVY)
    _set_para_space(team_label, 12, 4)

    team_headers = ["Role", "Name"]
    team_data = [
        ["Lead Faculty / IAC Director", director],
        ["Lead Student",                lead_student or "—"],
        ["Safety Student",              safety_stu or "—"],
        ["Additional Student(s)",       other_stu or "—"],
    ]
    _add_table(doc, team_data, headers=team_headers, col_widths=[3.0, 3.5])

    # Footer note on cover
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "University of Louisiana at Lafayette — MALT Industrial Assessment Center\n"
        "Funded by the U.S. Department of Energy, Office of Energy Efficiency and Renewable Energy"
    )
    _set_run_font(footer_run, "Calibri", 9, False, HEADER_GRAY)
    _set_para_space(footer_para, 24, 0)

    _add_page_break(doc)

    # ── DISCLAIMER ────────────────────────────────────────────────────────────
    _set_heading(doc, "DISCLAIMER", 1)
    _add_body_para(doc, DISCLAIMER_TEXT)
    _add_page_break(doc)

    # ── PREFACE ───────────────────────────────────────────────────────────────
    _set_heading(doc, "PREFACE", 1)
    _add_body_para(doc, PREFACE_TEXT)
    _add_page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    if include_toc:
        _add_toc(doc, ar_list)

    # ── SECTION 1: EXECUTIVE SUMMARY ─────────────────────────────────────────
    _set_heading(doc, "SECTION 1 — EXECUTIVE SUMMARY", 1)

    _set_heading(doc, "1.1  Annual Utility Usage and Cost", 2)
    rates = get_utility_rates_from_dict(session)

    ann_elec_total = (rates.get("ann_elec_cost", 0) +
                      rates.get("ann_demand_cost", 0) +
                      rates.get("ann_elec_fee", 0))
    ann_gas_total  = rates.get("ann_gas_cost", 0) + rates.get("ann_gas_fee", 0)
    ann_wtr_total  = (rates.get("ann_water_cost", 0) +
                      rates.get("ann_sewer_cost", 0) +
                      rates.get("ann_water_fee", 0))
    total_util = ann_elec_total + ann_gas_total + ann_wtr_total

    elec_rows_s = session.get("elec_rows", [])
    ann_kw_peak = max((r.get("kw", 0) for r in elec_rows_s), default=0)

    util_headers = ["Utility", "Consumption", "Peak Demand",
                    "Energy Cost", "Demand Cost", "Other Fees", "Total Cost"]
    util_data = []
    if rates.get("ann_kwh", 0) > 0:
        util_data.append([
            "Electricity",
            f"{rates['ann_kwh']:,.0f} kWh",
            f"{ann_kw_peak:,.0f} kW",
            f"${rates['ann_elec_cost']:,.0f}",
            f"${rates['ann_demand_cost']:,.0f}",
            f"${rates['ann_elec_fee']:,.0f}",
            f"${ann_elec_total:,.0f}",
        ])
    if rates.get("ann_mmbtu", 0) > 0:
        util_data.append([
            "Natural Gas",
            f"{rates['ann_mmbtu']:,.1f} MMBtu",
            "—",
            f"${rates['ann_gas_cost']:,.0f}",
            "—",
            f"${rates['ann_gas_fee']:,.0f}",
            f"${ann_gas_total:,.0f}",
        ])
    if rates.get("ann_tgal", 0) > 0:
        util_data.append([
            "Water/Sewer",
            f"{rates['ann_tgal']:,.3f} Tgal",
            "—",
            f"${rates['ann_water_cost'] + rates['ann_sewer_cost']:,.0f}",
            "—",
            f"${rates['ann_water_fee']:,.0f}",
            f"${ann_wtr_total:,.0f}",
        ])
    util_data.append(["TOTAL", "—", "—", "—", "—", "—", f"${total_util:,.0f}"])

    _add_table(doc, util_data, headers=util_headers,
               col_widths=[1.0, 1.0, 0.95, 0.9, 0.9, 0.9, 0.9])

    _set_heading(doc, "1.2  Summary of Assessment Recommendations", 2)
    if ar_list:
        ar_headers = ["AR #", "ARC Code", "Description",
                      "Resource Savings", "Cost Savings ($/yr)", "Impl. Cost ($)", "Payback (yr)"]
        ar_sum_data = []
        tot_cost = 0
        tot_impl = 0
        for ar in ar_list:
            res_parts = [
                f"{r['type']}: {r['savings']:,.0f} {r['unit']}"
                for r in ar.get("resources", []) if r.get("savings", 0) > 0
            ]
            pb = ar.get("payback", float("inf"))
            pb_str = f"{pb:.1f}" if pb not in (float("inf"), None) else "N/A"
            ar_sum_data.append([
                ar.get("ar_number", "—"),
                ar.get("arc_code", "—"),
                ar.get("title", "—"),
                " | ".join(res_parts) or "—",
                f"${ar.get('total_cost_savings', 0):,.0f}",
                f"${ar.get('implementation_cost', 0):,.0f}",
                pb_str,
            ])
            tot_cost += ar.get("total_cost_savings", 0)
            tot_impl += ar.get("implementation_cost", 0)

        avg_pb = tot_impl / tot_cost if tot_cost > 0 else float("inf")
        avg_pb_str = f"{avg_pb:.1f}" if avg_pb != float("inf") else "—"
        ar_sum_data.append(["—", "—", "TOTALS", "—",
                             f"${tot_cost:,.0f}", f"${tot_impl:,.0f}", avg_pb_str])
        _add_table(doc, ar_sum_data, headers=ar_headers,
                   col_widths=[0.55, 0.75, 1.65, 1.5, 0.95, 0.85, 0.75])
    else:
        _add_body_para(doc, "No assessment recommendations have been entered.")

    _add_page_break(doc)

    # ── SECTION 2: FACILITY BACKGROUND ────────────────────────────────────────
    _set_heading(doc, "SECTION 2 — GENERAL FACILITY BACKGROUND", 1)

    _set_heading(doc, "2.1  Facility Description", 2)
    _add_body_para(doc, session.get("facility_description", "") or "—")

    _set_heading(doc, "2.2  Process Description", 2)
    _add_body_para(doc, session.get("process_description", "") or "—")

    _set_heading(doc, "2.3  Best Practices", 2)
    best_practices = session.get("best_practices", [])
    bp_items = [b for b in best_practices if b]
    if bp_items:
        for bp in bp_items:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(bp)
            _set_run_font(run, "Calibri", 11, False, BLACK)
    else:
        _add_body_para(doc, "No best practices recorded.")

    _set_heading(doc, "2.4  Forms of Energy Usage", 2)
    elec_uses = session.get("elec_used_for", [])
    gas_uses  = session.get("gas_used_for", [])
    if elec_uses:
        para = doc.add_paragraph()
        r1 = para.add_run("Electricity is used for: ")
        _set_run_font(r1, "Calibri", 11, True, BLACK)
        r2 = para.add_run(", ".join(elec_uses))
        _set_run_font(r2, "Calibri", 11, False, BLACK)
    if gas_uses:
        para = doc.add_paragraph()
        r1 = para.add_run("Natural Gas is used for: ")
        _set_run_font(r1, "Calibri", 11, True, BLACK)
        r2 = para.add_run(", ".join(gas_uses))
        _set_run_font(r2, "Calibri", 11, False, BLACK)
    if not elec_uses and not gas_uses:
        _add_body_para(doc, "—")

    _set_heading(doc, "2.5  Major Energy Consuming Equipment", 2)
    eq_rows = session.get("equipment_rows", [])
    valid_eq = [r for r in eq_rows if r.get("equipment", "")]
    if valid_eq:
        eq_headers = ["#", "Equipment", "Specifications", "Qty / Capacity", "Energy Form"]
        eq_data = [
            [str(i + 1),
             r.get("equipment", ""),
             r.get("specs", ""),
             r.get("qty_capacity", ""),
             r.get("energy_form", "")]
            for i, r in enumerate(valid_eq)
        ]
        _add_table(doc, eq_data, headers=eq_headers,
                   col_widths=[0.35, 1.8, 2.0, 1.5, 1.0])
    else:
        _add_body_para(doc, "No equipment data entered.")

    _set_heading(doc, "2.6  Energy and Water Consumption with Cost", 2)

    # Electricity billing table
    elec_rows_data = session.get("elec_rows", [])
    if any(r.get("kwh", 0) > 0 for r in elec_rows_data):
        _add_body_para(doc, "Electricity Consumption and Cost", bold=True)
        e_headers = ["Month", "kWh", "Elec Cost", "kW Demand", "Demand Cost", "Fees", "Total"]
        e_data = []
        for r in elec_rows_data:
            if r.get("kwh", 0) > 0:
                row_total = r.get("elec_cost", 0) + r.get("demand_cost", 0) + r.get("fee", 0)
                e_data.append([
                    r.get("month", ""),
                    f"{r.get('kwh', 0):,.0f}",
                    f"${r.get('elec_cost', 0):,.2f}",
                    f"{r.get('kw', 0):,.0f}",
                    f"${r.get('demand_cost', 0):,.2f}",
                    f"${r.get('fee', 0):,.2f}",
                    f"${row_total:,.2f}",
                ])
        t_kwh = sum(r.get("kwh", 0) for r in elec_rows_data)
        t_ec  = sum(r.get("elec_cost", 0) for r in elec_rows_data)
        t_dc  = sum(r.get("demand_cost", 0) for r in elec_rows_data)
        t_ef  = sum(r.get("fee", 0) for r in elec_rows_data)
        e_data.append(["TOTAL", f"{t_kwh:,.0f}", f"${t_ec:,.2f}",
                        "—", f"${t_dc:,.2f}", f"${t_ef:,.2f}", f"${t_ec+t_dc+t_ef:,.2f}"])
        _add_table(doc, e_data, headers=e_headers,
                   col_widths=[0.65, 0.95, 0.95, 0.95, 0.95, 0.8, 0.85])

    # Natural gas billing table
    gas_rows_data = session.get("gas_rows", [])
    if any(r.get("mmbtu", 0) > 0 for r in gas_rows_data):
        _add_body_para(doc, "Natural Gas Consumption and Cost", bold=True)
        g_headers = ["Month", "MMBtu", "Gas Cost", "Fees", "Total"]
        g_data = []
        for r in gas_rows_data:
            if r.get("mmbtu", 0) > 0:
                row_total = r.get("cost", 0) + r.get("fee", 0)
                g_data.append([
                    r.get("month", ""),
                    f"{r.get('mmbtu', 0):,.1f}",
                    f"${r.get('cost', 0):,.2f}",
                    f"${r.get('fee', 0):,.2f}",
                    f"${row_total:,.2f}",
                ])
        t_m  = sum(r.get("mmbtu", 0) for r in gas_rows_data)
        t_gc = sum(r.get("cost", 0) for r in gas_rows_data)
        t_gf = sum(r.get("fee", 0) for r in gas_rows_data)
        g_data.append(["TOTAL", f"{t_m:,.1f}", f"${t_gc:,.2f}",
                        f"${t_gf:,.2f}", f"${t_gc+t_gf:,.2f}"])
        _add_table(doc, g_data, headers=g_headers,
                   col_widths=[0.75, 1.1, 1.2, 1.1, 1.1])

    # Water billing table
    water_rows_data = session.get("water_rows", [])
    if any(r.get("tgal", 0) > 0 for r in water_rows_data):
        _add_body_para(doc, "Water and Sewer Consumption and Cost", bold=True)
        w_headers = ["Month", "Tgal", "Water Cost", "Sewer Cost", "Fees", "Total"]
        w_data = []
        for r in water_rows_data:
            if r.get("tgal", 0) > 0:
                row_total = r.get("water_cost", 0) + r.get("sewer_cost", 0) + r.get("fee", 0)
                w_data.append([
                    r.get("month", ""),
                    f"{r.get('tgal', 0):,.3f}",
                    f"${r.get('water_cost', 0):,.2f}",
                    f"${r.get('sewer_cost', 0):,.2f}",
                    f"${r.get('fee', 0):,.2f}",
                    f"${row_total:,.2f}",
                ])
        t_tg = sum(r.get("tgal", 0) for r in water_rows_data)
        t_wc = sum(r.get("water_cost", 0) for r in water_rows_data)
        t_sc = sum(r.get("sewer_cost", 0) for r in water_rows_data)
        t_wf = sum(r.get("fee", 0) for r in water_rows_data)
        w_data.append(["TOTAL", f"{t_tg:,.3f}", f"${t_wc:,.2f}",
                        f"${t_sc:,.2f}", f"${t_wf:,.2f}", f"${t_wc+t_sc+t_wf:,.2f}"])
        _add_table(doc, w_data, headers=w_headers,
                   col_widths=[0.75, 0.85, 1.0, 1.0, 0.85, 0.9])

    _add_page_break(doc)

    # ── SECTION 3: ASSESSMENT RECOMMENDATIONS ────────────────────────────────
    _set_heading(doc, "SECTION 3 — ENERGY ASSESSMENT RECOMMENDATIONS", 1)
    if ar_list:
        for idx, ar in enumerate(ar_list):
            _add_ar_section(doc, ar, idx + 1)
    else:
        _add_body_para(doc, "No assessment recommendations have been entered.")

    _add_page_break(doc)

    # ── SECTION 4: CYBERSECURITY ──────────────────────────────────────────────
    if include_cyber:
        _set_heading(doc, "SECTION 4 — CYBERSECURITY ASSESSMENT", 1)
        _add_body_para(doc, CYBER_TEXT)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
